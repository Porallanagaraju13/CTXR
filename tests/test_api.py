import pytest
from fastapi.testclient import TestClient
from contextforge.api.app import app

client = TestClient(app)

def test_health_check_endpoint():
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json()["status"] == "healthy"
    assert "supported_formats" in response.json()

def test_normalize_empty_payload_rejected():
    response = client.post("/normalize")
    assert response.status_code == 422 # Unprocessable upload

def test_normalize_unsupported_format_rejected():
    response = client.post(
        "/normalize",
        files={"file": ("malicious.exe", b"binary-data", "application/octet-stream")},
        data={"use_ai": "false"}
    )
    assert response.status_code == 400
    assert "Unsupported format" in response.json()["detail"]

def test_normalize_valid_document_with_configurations(tmp_path):
    import docx
    doc_path = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_heading("ContextForge Architecture", level=1)
    doc.add_paragraph("This is an Organization called ContextForge Corp working on FastAPI since 2026-06-02.")
    doc.save(str(doc_path))
    
    with open(doc_path, "rb") as f:
        doc_bytes = f.read()
        
    response = client.post(
        "/normalize",
        files={"file": ("sample.docx", doc_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={
            "use_ai": "false",
            "chunk_size": "200",
            "target_model": "gpt-4o",
            "extract_entities": "true",
            "deduplicate_boilerplate": "false"
        }
    )
    assert response.status_code == 200
    res_data = response.json()
    assert res_data["document_name"] == "sample.docx"
    assert "metrics" in res_data
    assert "dollars_saved" in res_data["metrics"]
    assert res_data["metadata"]["target_model"] == "gpt-4o"
    # Verify entity extraction worked
    assert len(res_data["entities"]) > 0
