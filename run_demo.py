import os
from pathlib import Path
import docx
from rich.console import Console
from rich.panel import Panel
from rich.table import Table

# Import ContextForge Engine components
from contextforge.adapters.local_parsers import LocalDocumentParser
from contextforge.services.compression import TokenCompressor
from contextforge.services.chunking import SemanticChunker, EntityExtractor
from contextforge.analytics import compute_savings

console = Console()

def create_demo_document(filepath: Path):
    """Generates a bloated sample Word document for structural optimization testing."""
    doc = docx.Document()
    doc.add_heading("ContextForge System Specifications", level=1)
    
    doc.add_paragraph("ContextForge   is   a   high-performance   normalization   engine.")
    doc.add_paragraph("Page 1 of 12") # Boilerplate header
    doc.add_paragraph("Copyright © 2026 ContextForge Corp. All rights reserved.") # Boilerplate footer
    
    doc.add_heading("1. Local Extraction Modules", level=2)
    doc.add_paragraph("The engineering team at ContextForge Corp released local parsers on 15 March 2026.")
    
    # Introduce duplicate layout OCR block
    doc.add_paragraph("This is a duplicated paragraph that simulated OCR line overlap detection errors.")
    doc.add_paragraph("This is a duplicated paragraph that simulated OCR line overlap detection errors.")
    
    # Add a table
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Engine Name"
    table.cell(0, 1).text = "Parsing Speed"
    table.cell(0, 2).text = "Token Savings"
    
    table.cell(1, 0).text = "Local Parser"
    table.cell(1, 1).text = "950 pages/sec"
    table.cell(1, 2).text = "42%"
    
    table.cell(2, 0).text = "AI Visual Parser"
    table.cell(2, 1).text = "45 pages/sec"
    table.cell(2, 2).text = "74%"
    
    doc.save(str(filepath))

def run_once_demo():
    console.print(Panel.fit(
        "[bold cyan]ContextForge Core Engine - Asynchronous Single-Run execution[/bold cyan]\n"
        "[dim]Processing: heterogeneous input document -> intelligent normalization -> pre-computed context package[/dim]",
        border_style="cyan"
    ))
    
    demo_file = Path("demo_sample.docx")
    
    # Step 1: Generate dynamic bloated sample document
    console.print("[yellow]Creating bloated Word document with duplicates and copyright headers...[/yellow]")
    create_demo_document(demo_file)
    
    # Step 2: Local document parsing
    console.print("[yellow]Parsing structure using LocalDocumentParser...[/yellow]")
    pages = LocalDocumentParser.parse_docx(demo_file)
    raw_content = "\n\n".join([p.text_content for p in pages])
    
    # Step 3: Optimization & Compression
    console.print("[yellow]Pruning boilerplate metadata and compressing spacing...[/yellow]")
    clean_markdown = TokenCompressor.strip_metadata_bloat(raw_content)
    clean_markdown = TokenCompressor.compress_syntax(clean_markdown)
    
    # Step 4: Chunks & Entity Tagging
    console.print("[yellow]Extracting semantic entities and segmenting chunks...[/yellow]")
    chunks = SemanticChunker.chunk_markdown(clean_markdown, max_chunk_words=100)
    entities = EntityExtractor.extract_entities(clean_markdown)
    
    # Step 5: Compute Savings Analytics
    raw_bytes = demo_file.stat().st_size
    raw_tokens_estimate = raw_bytes // 2 # byte to token proxy
    stats = compute_savings(raw_tokens_estimate, clean_markdown)
    
    # Save the output package
    output_markdown = Path("demo_normalized.md")
    output_markdown.write_text(clean_markdown, encoding="utf-8")
    
    # Step 6: Visual Scorecard Output
    console.print("\n[bold green][SUCCESS] Normalization pipeline executed successfully![/bold green]")
    console.print(f"Normalized GFM Markdown written to: [yellow]{output_markdown}[/yellow]\n")
    
    # Render GFM Content Output
    console.print(Panel(
        f"[bold white]Pre-computed GFM Markdown Package:[/bold white]\n\n{clean_markdown}",
        title="gfm_normalized_content",
        border_style="green"
    ))
    
    # Render Metadata and Entities
    console.print("\n[bold cyan]Indexed Entities:[/bold cyan]")
    for idx, ent in enumerate(entities, 1):
        console.print(f" {idx}. [bold white]{ent.name}[/bold white] ({ent.category})")
        
    # Render Token Savings Panel
    table = Table(title="ContextForge Token Savings Scorecard", show_header=True, header_style="bold magenta")
    table.add_column("Optimization Metric", style="cyan")
    table.add_column("Value", style="bold white")
    
    table.add_row("Estimated Raw Footprint", f"{stats['raw_tokens']:,} tokens")
    table.add_row("Pre-computed Context Package", f"{stats['markdown_tokens']:,} tokens")
    table.add_row("Total Context Footprint Saved", f"[bold green]{stats['tokens_saved']:,} tokens[/bold green]")
    table.add_row("Context Footprint Reduction", f"[bold green]{stats['savings_percentage']}%[/bold green]")
    
    console.print("\n")
    console.print(table)
    
    # Clean up demo files
    if demo_file.exists():
        os.unlink(demo_file)
        
if __name__ == "__main__":
    run_once_demo()
